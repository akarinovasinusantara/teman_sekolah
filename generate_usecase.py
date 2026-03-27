from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p

doc = Document()

add_heading(doc, "Teman Sekolah - Use Case Diagram & Spesifikasi", level=1, color=(21, 101, 192))
add_paragraph(doc, "Dokumen ini berisi spesifikasi Use Case untuk sistem manajemen sekolah terpadu Teman Sekolah.", bold=True)

add_heading(doc, "1. Daftar Aktor", level=2)
actors = [
    ("Super Admin (Yayasan/Pengembang)", "Aktor yang mengelola keseluruhan data sekolah, laporan keuangan terpusat, dan mengatur langganan."),
    ("Tata Usaha (TU)", "Aktor yang menangani operasional harian sekolah, seperti PPDB, keuangan SPP, jadwal, dan surat menyurat."),
    ("Guru", "Aktor yang berinteraksi dalam proses akademik, seperti absensi kelas, nilai, e-rapor, dan pengumuman."),
    ("Orang Tua / Siswa", "Aktor yang bertindak sebagai penerima informasi terkait rekam akademik, kehadiran, dan tagihan keuangan.")
]

for actor, desc in actors:
    add_paragraph(doc, actor, bold=True)
    add_bullet(doc, desc)

add_heading(doc, "2. Daftar Use Case", level=2)

use_cases = [
    ("Super Admin", [
        ("Melihat Dashboard", "Mengakses ringkasan statistik sekolah dan pengguna."),
        ("Kelola Data Sekolah", "Menambah, mengubah, dan menghapus (CRUD) data sekolah."),
        ("Lihat Laporan Keuangan", "Memantau laporan pemasukan masuk dan tunggakan total SPP."),
        ("Monitoring Kinerja Guru", "Memonitor jurnal mengajar dan evaluasi kinerja guru."),
        ("Pengaturan & Langganan", "Mengatur biaya langganan aplikasi.")
    ]),
    ("Tata Usaha (TU)", [
        ("Melihat Dashboard", "Mengecek ringkasan statistik sekolah harian."),
        ("Kelola PPDB", "Memantau, menyetujui, dan mengatur pendaftaran siswa baru secara online."),
        ("Kelola Keuangan SPP", "Membuat tagihan, serta mengonfirmasi pembayaran masuk dari siswa."),
        ("Atur Jadwal Pelajaran", "Menyusun jadwal pelajaran per kelas dan per guru."),
        ("Kelola Surat Menyurat", "Mencetak dan membuat surat edaran resmi serta ijazah.")
    ]),
    ("Guru", [
        ("Melihat Dashboard", "Melihat kelas dan jadwal mengajar."),
        ("Input Absensi Siswa", "Merekam presensi/kehadiran siswa setiap pertemuan kelas berlangsung."),
        ("Input Nilai & E-Rapor", "Memasukkan nilai (harian, UTS, UAS) serta mengisi deskripsi rapor."),
        ("Kirim Pengumuman Kelas", "Memberikan tugas atau jadwal informasi terhadap kelas yang diampu.")
    ]),
    ("Orang Tua / Siswa", [
        ("Melihat Dashboard", "Melihat info pribadi dan rangkuman aktivitas."),
        ("Bayar Tagihan & SPP", "Melihat tagihan berjalan lalu melakukan pembayaran mandiri (QRIS/VA)."),
        ("Monitoring Kehadiran Siswa", "Mengecek kehadiran secara aktual (real-time)."),
        ("Lihat Hasil Belajar & Rapor", "Mengakses rapor digital dan rincian nilai ujian."),
        ("Terima Informasi", "Mendapatkan notifikasi dan surat dari pihak sekolah/guru.")
    ])
]

for role, ucs in use_cases:
    add_heading(doc, f"Role: {role}", level=3)
    for uc, desc in ucs:
        add_paragraph(doc, uc, bold=True)
        add_bullet(doc, desc)

add_heading(doc, "3. Mermaid Code (Use Case Diagram)", level=2)
add_paragraph(doc, "Salin kode Mermaid di bawah ini ke Markdown/Mermaid Live Editor untuk melihat bentuk visual diagram:")

mermaid_code = """flowchart LR
    classDef actor fill:#f9f9f9,stroke:#333,stroke-width:2px;
    classDef usecase fill:#e1f5fe,stroke:#03a9f4,stroke-width:2px,rx:20,ry:20;
    classDef sys fill:#fff,stroke:#666,stroke-width:2px,stroke-dasharray: 5 5;

    SA(Super Admin):::actor
    TU(Tata Usaha):::actor
    G(Guru):::actor
    O(Orang Tua/Siswa):::actor

    subgraph System["Teman Sekolah"]
        direction TB
        UC_Dash([Melihat Dashboard]):::usecase
        UC_Auth([Login / Autentikasi]):::usecase
        
        UC_ManSek([Kelola Data Sekolah]):::usecase
        UC_PPDB([Kelola PPDB]):::usecase
        UC_Absen([Input Absensi Siswa]):::usecase
        UC_ByrSPP([Bayar Tagihan & SPP]):::usecase
    end

    SA --> UC_Auth
    TU --> UC_Auth
    G --> UC_Auth
    O --> UC_Auth
    
    SA --> UC_ManSek
    TU --> UC_PPDB
    G --> UC_Absen
    O --> UC_ByrSPP
    class System sys;
"""

p = doc.add_paragraph()
r = p.add_run(mermaid_code)
r.font.name = 'Courier New'
r.font.size = Pt(9)

output_path = r"d:\Project\Teman_Sekolah\UseCase_Teman_Sekolah.docx"
doc.save(output_path)
print(f"File Use Case Word berhasil dibuat: {output_path}")
