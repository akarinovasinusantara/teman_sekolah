
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─────────────────────────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────────────────────────

def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1565C0')
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size  = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E3F2FD')
                tcPr.append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)

    return table

def add_section_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1565C0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

# ─────────────────────────────────────────────────────────────────────────────
# Page setup
# ─────────────────────────────────────────────────────────────────────────────
set_page_margins(doc)

# ─────────────────────────────────────────────────────────────────────────────
# Cover Page
# ─────────────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run("PRODUCT REQUIREMENTS DOCUMENT")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(21, 101, 192)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_sub.add_run("Teman Sekolah")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(13, 71, 161)

cover_desc = doc.add_paragraph()
cover_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_desc.add_run("Sistem Manajemen Sekolah Terpadu")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(66, 66, 66)

doc.add_paragraph()
doc.add_paragraph()

meta_table = doc.add_table(rows=6, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Versi Dokumen", "1.0"),
    ("Tanggal",       "25 Maret 2026"),
    ("Status",        "Draft"),
    ("Tim",           "Teman Sekolah Development Team"),
    ("Platform",      "Web Application (React + Node.js)"),
    ("Bahasa",        "Indonesia"),
]
for i, (label, value) in enumerate(meta_data):
    row = meta_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(21, 101, 192)
            run.font.size = Pt(11)
    for para in row.cells[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. Pendahuluan
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Pendahuluan", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "1.1 Latar Belakang", level=2)
add_paragraph(doc,
    "Pengelolaan administrasi sekolah secara manual masih menjadi tantangan besar bagi banyak "
    "institusi pendidikan di Indonesia. Proses pencatatan absensi, nilai, pembayaran SPP, dan "
    "komunikasi antara sekolah dengan orang tua siswa seringkali dilakukan secara terpisah, tidak "
    "terintegrasi, dan rentan terhadap kesalahan. Hal ini mengakibatkan inefisiensi operasional, "
    "keterlambatan informasi, dan menurunnya kualitas pelayanan pendidikan.")
add_paragraph(doc,
    "Teman Sekolah hadir sebagai solusi digital terpadu yang mengintegrasikan seluruh proses "
    "manajemen sekolah dalam satu platform. Dengan memanfaatkan teknologi web modern, aplikasi ini "
    "memungkinkan seluruh pemangku kepentingan - mulai dari yayasan, staf administrasi, guru, "
    "hingga orang tua dan siswa - untuk mengakses informasi dan menjalankan tugas mereka secara "
    "efisien dan real-time.")

add_heading(doc, "1.2 Tujuan Dokumen", level=2)
add_paragraph(doc, "Dokumen PRD (Product Requirements Document) ini bertujuan untuk:")
add_bullet(doc, "Mendefinisikan visi, tujuan, dan ruang lingkup produk Teman Sekolah.")
add_bullet(doc, "Mendeskripsikan persyaratan fungsional dan non-fungsional secara lengkap.")
add_bullet(doc, "Menjadi panduan bagi tim pengembang, desainer, dan pemangku kepentingan.")
add_bullet(doc, "Menjadi acuan dalam pengujian dan validasi produk.")

add_heading(doc, "1.3 Ruang Lingkup", level=2)
add_paragraph(doc, "Teman Sekolah adalah aplikasi web manajemen sekolah berbasis multi-role yang mencakup:")
add_bullet(doc, "Manajemen pengguna dengan 5 peran berbeda (Super Admin, TU, Guru, Orang Tua, Siswa).")
add_bullet(doc, "Modul PPDB (Penerimaan Peserta Didik Baru) online.")
add_bullet(doc, "Manajemen absensi, penilaian, dan e-rapor digital.")
add_bullet(doc, "Sistem keuangan SPP terintegrasi dengan notifikasi pembayaran.")
add_bullet(doc, "Jadwal pelajaran, pengumuman kelas, dan surat edaran digital.")
add_bullet(doc, "Dashboard analitik untuk setiap peran pengguna.")

add_heading(doc, "1.4 Definisi dan Singkatan", level=2)
add_table(doc, ["Istilah", "Definisi"], [
    ["PRD",    "Product Requirements Document - dokumen persyaratan produk"],
    ["PPDB",   "Penerimaan Peserta Didik Baru - proses pendaftaran siswa baru"],
    ["SPP",    "Sumbangan Pembinaan Pendidikan - iuran bulanan siswa"],
    ["TU",     "Tata Usaha - staf administrasi sekolah"],
    ["JWT",    "JSON Web Token - standar autentikasi berbasis token"],
    ["MUI",    "Material-UI - library komponen React berbasis Material Design"],
    ["ORM",    "Object Relational Mapping - abstraksi database menggunakan Sequelize"],
    ["NPSN",   "Nomor Pokok Sekolah Nasional - identitas resmi sekolah"],
    ["NIS",    "Nomor Induk Siswa - identitas unik setiap siswa"],
    ["NIP",    "Nomor Induk Pegawai - identitas unik setiap guru/pegawai"],
    ["API",    "Application Programming Interface - antarmuka komunikasi antar sistem"],
    ["SaaS",   "Software as a Service - model layanan perangkat lunak berbasis langganan"],
], col_widths=[3.5, 13.0])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 2. Visi & Misi Produk
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Visi & Misi Produk", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "2.1 Visi Produk", level=2)
p = doc.add_paragraph()
run = p.add_run(
    '"Menjadi platform manajemen sekolah terdepan di Indonesia yang menghadirkan ekosistem '
    'pendidikan digital terintegrasi, efisien, dan dapat diakses oleh semua lapisan pemangku kepentingan."')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(21, 101, 192)
p.paragraph_format.left_indent = Cm(1)

add_heading(doc, "2.2 Misi Produk", level=2)
add_bullet(doc, "Mendigitalkan dan mengotomasi proses administrasi sekolah secara menyeluruh.")
add_bullet(doc, "Meningkatkan transparansi informasi antara sekolah, guru, dan orang tua.")
add_bullet(doc, "Menyediakan analitik real-time untuk pengambilan keputusan yang lebih baik.")
add_bullet(doc, "Membantu sekolah meningkatkan efisiensi operasional dan kualitas pendidikan.")
add_bullet(doc, "Menyediakan platform SaaS yang skalabel dan mudah diadopsi oleh berbagai sekolah.")

add_heading(doc, "2.3 Proposisi Nilai (Value Proposition)", level=2)
add_table(doc, ["Segmen Pengguna", "Nilai yang Diberikan"], [
    ["Super Admin / Yayasan", "Visibilitas penuh terhadap kinerja sekolah, laporan keuangan terpusat, monitoring kinerja guru"],
    ["Staf TU",               "Otomasi proses PPDB, pengelolaan SPP digital, manajemen jadwal & surat"],
    ["Guru",                  "Absensi digital, input nilai mudah, e-rapor otomatis, komunikasi kelas efektif"],
    ["Orang Tua",             "Monitoring anak real-time, pembayaran SPP online, terima informasi sekolah"],
    ["Siswa",                 "Akses nilai & rapor digital, lihat jadwal & kehadiran secara mandiri"],
], col_widths=[4.5, 12.0])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 3. Target Pengguna
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Target Pengguna", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "3.1 Segmen Pasar", level=2)
add_paragraph(doc,
    "Teman Sekolah ditujukan untuk sekolah-sekolah formal di Indonesia, mulai dari jenjang "
    "SD, SMP, hingga SMA/SMK yang ingin bertransisi ke sistem manajemen berbasis digital.")

add_heading(doc, "3.2 User Persona", level=2)

personas = [
    ("Super Admin - Yayasan/Pengembang",
     "Pengelola yayasan atau pengembang aplikasi yang memiliki akses penuh ke seluruh sistem.",
     ["Memonitor kinerja seluruh sekolah dalam jaringan yayasan",
      "Melihat laporan keuangan terpusat (pemasukan & tunggakan SPP)",
      "Memantau kinerja mengajar guru melalui jurnal mengajar",
      "Mengatur biaya langganan aplikasi per sekolah",
      "Menambah dan mengelola data sekolah baru"]),
    ("Staf TU - Tata Usaha",
     "Pegawai administrasi yang menangani operasional harian sekolah.",
     ["Memproses pendaftaran siswa baru (PPDB) secara digital",
      "Mengelola tagihan dan pembayaran SPP siswa",
      "Mengatur jadwal pelajaran per kelas dan per semester",
      "Mencetak dan mengirimkan surat edaran digital",
      "Membuat dan mengelola akun pengguna untuk guru dan orang tua"]),
    ("Guru",
     "Pengajar yang bertanggung jawab atas proses pembelajaran dan penilaian siswa.",
     ["Input kehadiran siswa secara digital per pertemuan",
      "Input nilai harian, UTS, dan UAS dengan mudah",
      "Mengisi e-rapor dan deskripsi perkembangan siswa",
      "Mengirimkan pengumuman dan tugas ke kelas yang diampu",
      "Melihat jadwal mengajar dan data kelas"]),
    ("Orang Tua",
     "Wali siswa yang ingin memantau perkembangan anak secara real-time.",
     ["Melihat kehadiran anak secara real-time",
      "Melihat nilai dan rapor digital anak",
      "Membayar SPP secara online (QRIS/Virtual Account)",
      "Menerima surat edaran dan pengumuman dari sekolah",
      "Mendapatkan notifikasi bila ada informasi penting"]),
    ("Siswa",
     "Peserta didik yang menggunakan platform untuk mengakses informasi akademik.",
     ["Melihat jadwal pelajaran secara digital",
      "Memantau rekap kehadiran pribadi",
      "Mengakses nilai dan rapor digital",
      "Melihat tagihan SPP dan status pembayaran",
      "Menerima pengumuman dari guru dan sekolah"]),
]

for judul, deskripsi, kebutuhan in personas:
    add_heading(doc, judul, level=3)
    add_paragraph(doc, deskripsi, italic=True)
    add_paragraph(doc, "Kebutuhan utama:", bold=True, space_after=2)
    for need in kebutuhan:
        add_bullet(doc, need)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 4. Persyaratan Fungsional
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. Persyaratan Fungsional", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "4.1 Modul Autentikasi & Manajemen Pengguna", level=2)
add_table(doc, ["ID", "Fitur", "Deskripsi", "Role"], [
    ["F-001", "Login",               "User dapat login menggunakan username & password",                        "Semua"],
    ["F-002", "JWT Token",           "Sistem menghasilkan token JWT dengan masa berlaku 7 hari",               "Sistem"],
    ["F-003", "Role-Based Access",   "Setiap user hanya dapat mengakses menu dan fitur sesuai perannya",       "Semua"],
    ["F-004", "Manajemen User",      "CRUD pengguna (tambah, lihat, edit, hapus)",                             "Super Admin, TU"],
    ["F-005", "Ganti Password",      "User dapat mengubah password pribadi",                                   "Semua"],
    ["F-006", "Logout",              "User dapat mengakhiri sesi dengan aman",                                 "Semua"],
], col_widths=[1.5, 3.5, 9.0, 2.5])
doc.add_paragraph()

add_heading(doc, "4.2 Modul Super Admin", level=2)
add_table(doc, ["ID", "Fitur", "Deskripsi"], [
    ["F-010", "Dashboard Super Admin",   "Ringkasan statistik: jumlah sekolah, guru, siswa, pemasukan, dan tunggakan"],
    ["F-011", "Manajemen Sekolah",       "CRUD data sekolah: nama, NPSN, alamat, status aktif, biaya langganan"],
    ["F-012", "Laporan Keuangan",        "Laporan pemasukan SPP total & per sekolah, tunggakan, dan grafik tren"],
    ["F-013", "Monitoring Kinerja Guru", "Lihat jumlah jurnal mengajar per guru, rata-rata kinerja, dan peringkat"],
    ["F-014", "Pengaturan Aplikasi",     "Atur biaya langganan, fitur aktif/nonaktif per sekolah"],
], col_widths=[1.5, 4.0, 11.0])
doc.add_paragraph()

add_heading(doc, "4.3 Modul Staf TU (Tata Usaha)", level=2)
add_table(doc, ["ID", "Fitur", "Deskripsi"], [
    ["F-020", "Dashboard TU",       "Ringkasan: total siswa aktif, guru, tunggakan SPP, jadwal hari ini"],
    ["F-021", "PPDB Online",        "Formulir pendaftaran publik, kelola status (pending/diterima/ditolak), cetak bukti"],
    ["F-022", "Keuangan SPP",       "Buat tagihan bulanan, konfirmasi pembayaran, ekspor laporan tunggakan"],
    ["F-023", "Jadwal Pelajaran",   "Buat & edit jadwal mata pelajaran per kelas/semester, deteksi konflik"],
    ["F-024", "Surat Menyurat",     "Buat surat edaran digital, cetak ijazah/SKHUN, distribusi ke orang tua"],
    ["F-025", "Data Siswa",         "CRUD data siswa: NIS, NISN, nama, kelas, data orang tua"],
    ["F-026", "Data Guru",          "CRUD data guru: NIP, nama, bidang studi, pendidikan terakhir"],
    ["F-027", "Data Kelas",         "CRUD kelas: nama kelas, wali kelas, tahun ajaran, kapasitas"],
], col_widths=[1.5, 4.0, 11.0])
doc.add_paragraph()

add_heading(doc, "4.4 Modul Guru", level=2)
add_table(doc, ["ID", "Fitur", "Deskripsi"], [
    ["F-030", "Dashboard Guru",     "Ringkasan: jumlah kelas, jadwal hari ini, siswa yang diabsen, tugas pending"],
    ["F-031", "Absensi Siswa",      "Input kehadiran per pertemuan (Hadir/Izin/Sakit/Alfa), rekap per bulan"],
    ["F-032", "Input Nilai",        "Input nilai harian, UTS, UAS per mata pelajaran per siswa"],
    ["F-033", "E-Rapor",            "Isi deskripsi perkembangan siswa, hitung nilai akhir, cetak rapor digital"],
    ["F-034", "Pengumuman Kelas",   "Kirim pengumuman ke kelas tertentu, atur prioritas dan tanggal kadaluarsa"],
    ["F-035", "Jurnal Mengajar",    "Catat aktivitas mengajar harian sebagai dokumentasi kinerja"],
], col_widths=[1.5, 4.0, 11.0])
doc.add_paragraph()

add_heading(doc, "4.5 Modul Orang Tua & Siswa", level=2)
add_table(doc, ["ID", "Fitur", "Deskripsi"], [
    ["F-040", "Dashboard",              "Ringkasan kehadiran, nilai terbaru, tagihan SPP, pengumuman terbaru"],
    ["F-041", "Pembayaran SPP",         "Lihat riwayat tagihan, status pembayaran, bayar via QRIS/Virtual Account"],
    ["F-042", "Monitoring Kehadiran",   "Lihat rekap kehadiran bulanan & tahunan secara real-time"],
    ["F-043", "Hasil Belajar",          "Lihat nilai per mata pelajaran, UTS, UAS, dan rapor digital"],
    ["F-044", "Informasi & Surat",      "Terima surat edaran sekolah, pengumuman guru, dan notifikasi penting"],
], col_widths=[1.5, 4.0, 11.0])
doc.add_paragraph()

add_heading(doc, "4.6 API Endpoints Utama", level=2)
add_table(doc, ["Module", "Method", "Endpoint", "Deskripsi", "Auth"], [
    ["Auth",     "POST", "/api/auth/login",           "Login dan dapatkan JWT token",                "Public"],
    ["Auth",     "GET",  "/api/auth/me",              "Profil user yang sedang login",               "Token"],
    ["Auth",     "POST", "/api/auth/register",        "Daftarkan user baru",                         "Super Admin"],
    ["Users",    "GET",  "/api/users",                "Daftar semua pengguna",                       "Admin, TU"],
    ["Users",    "PUT",  "/api/users/:id",            "Update data pengguna",                        "Admin, TU"],
    ["Sekolah",  "GET",  "/api/sekolah",              "Daftar sekolah",                              "Super Admin"],
    ["Sekolah",  "POST", "/api/sekolah",              "Tambah sekolah baru",                         "Super Admin"],
    ["Siswa",    "GET",  "/api/siswa",                "Daftar siswa",                                "Token"],
    ["Siswa",    "POST", "/api/siswa",                "Tambah siswa baru",                           "TU, Admin"],
    ["Guru",     "GET",  "/api/guru",                 "Daftar guru",                                 "Token"],
    ["Absensi",  "GET",  "/api/absensi",              "List data absensi",                           "Token"],
    ["Absensi",  "POST", "/api/absensi",              "Input absensi baru",                          "Guru, TU"],
    ["Nilai",    "GET",  "/api/nilai/siswa/:id",      "Nilai per siswa",                             "Token"],
    ["Nilai",    "POST", "/api/nilai",                "Input nilai siswa",                           "Guru, TU"],
    ["Keuangan", "GET",  "/api/keuangan/summary/total","Ringkasan keuangan",                         "Admin, TU"],
    ["Keuangan", "POST", "/api/keuangan",             "Buat tagihan/transaksi",                      "TU"],
    ["PPDB",     "POST", "/api/ppdb",                 "Formulir pendaftaran (publik)",               "Public"],
    ["PPDB",     "PUT",  "/api/ppdb/:id",             "Update status pendaftaran",                   "TU"],
], col_widths=[2.0, 1.5, 5.5, 5.5, 2.0])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 5. Persyaratan Non-Fungsional
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Persyaratan Non-Fungsional", level=1, color=(21, 101, 192))
add_section_line(doc)

add_table(doc, ["Kategori", "Persyaratan", "Target / Kriteria"], [
    ["Performa",       "Waktu respons API",              "<= 500ms untuk 95% request"],
    ["Performa",       "Waktu muat halaman",             "<= 2 detik (First Contentful Paint)"],
    ["Performa",       "Kapasitas concurrent users",     "Mendukung 500 pengguna simultan"],
    ["Keamanan",       "Autentikasi",                    "JWT dengan enkripsi HS256, expire 7 hari"],
    ["Keamanan",       "Password hashing",               "Bcrypt dengan salt rounds >= 10"],
    ["Keamanan",       "HTTPS",                          "Wajib menggunakan HTTPS di lingkungan produksi"],
    ["Keamanan",       "Rate limiting",                  "Maksimal 100 request/menit per IP"],
    ["Keamanan",       "SQL Injection & XSS",            "Validasi dan sanitasi input di semua endpoint"],
    ["Skalabilitas",   "Arsitektur",                     "Modular dan siap untuk horizontal scaling"],
    ["Skalabilitas",   "Database",                       "SQLite untuk dev, MySQL/PostgreSQL untuk produksi"],
    ["Ketersediaan",   "Uptime",                         "99.5% uptime per bulan (SLA)"],
    ["Ketersediaan",   "Backup database",                "Otomatis setiap 24 jam"],
    ["Responsivitas",  "Mobile-friendly",                "Responsif di layar 320px - 2560px"],
    ["Responsivitas",  "Browser support",                "Chrome, Firefox, Safari, Edge (versi terbaru)"],
    ["Maintainability","Kode",                           "Modular, terdokumentasi, mengikuti standar ESLint"],
    ["Maintainability","API Documentation",              "Swagger/OpenAPI tersedia di /api/docs"],
], col_widths=[3.5, 5.5, 7.5])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 6. Arsitektur Teknis
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Arsitektur Teknis", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "6.1 Stack Teknologi", level=2)
add_table(doc, ["Layer", "Teknologi", "Versi", "Fungsi"], [
    ["Frontend",   "React",            "18.x",   "UI Library utama"],
    ["Frontend",   "Vite",             "5.x",    "Build tool & dev server"],
    ["Frontend",   "Material-UI",      "5.x",    "Library komponen UI"],
    ["Frontend",   "React Router DOM", "6.x",    "Client-side routing"],
    ["Frontend",   "Axios",            "1.x",    "HTTP client untuk API calls"],
    ["Backend",    "Node.js",          "18+",    "Runtime environment"],
    ["Backend",    "Express.js",       "4.x",    "Web framework REST API"],
    ["Backend",    "Sequelize",        "6.x",    "ORM untuk database"],
    ["Backend",    "JWT",              "9.x",    "Token autentikasi"],
    ["Backend",    "Bcryptjs",         "2.x",    "Hashing password"],
    ["Database",   "SQLite",           "3.x",    "Database development"],
    ["Database",   "MySQL/PostgreSQL", "-",      "Database produksi (rekomendasi)"],
], col_widths=[2.5, 3.5, 2.0, 8.5])
doc.add_paragraph()

add_heading(doc, "6.2 Arsitektur Sistem", level=2)
add_paragraph(doc, "Teman Sekolah menggunakan arsitektur client-server dengan pemisahan yang jelas antara frontend dan backend:")
add_bullet(doc, "Frontend: Single Page Application (SPA) berbasis React yang berkomunikasi dengan backend melalui RESTful API.")
add_bullet(doc, "Backend: RESTful API menggunakan Express.js dengan middleware autentikasi JWT.")
add_bullet(doc, "Database: ORM Sequelize yang mendukung SQLite (pengembangan) dan MySQL/PostgreSQL (produksi).")
add_bullet(doc, "Autentikasi: Stateless JWT yang disimpan di localStorage frontend.")

add_heading(doc, "6.3 Model Database", level=2)
add_table(doc, ["Model", "Field Utama", "Deskripsi"], [
    ["User",           "user_id, username, password, role, nama_lengkap", "Pengguna sistem dengan 5 peran"],
    ["Sekolah",        "nama, npsn, alamat, status, biaya_langganan",     "Data institusi sekolah"],
    ["Siswa",          "nis, nisn, nama_lengkap, kelas_id, orang_tua",    "Data peserta didik"],
    ["Guru",           "nip, nama_lengkap, pendidikan, total_jurnal",     "Data tenaga pengajar"],
    ["Kelas",          "nama, tingkat, wali_kelas_id, tahun_ajaran",      "Data rombongan belajar"],
    ["MataPelajaran",  "kode, nama, kelompok, kkm",                       "Daftar mata pelajaran"],
    ["JadwalPelajaran","kelas_id, mapel_id, guru_id, hari, jam",          "Jadwal mengajar per kelas"],
    ["Absensi",        "siswa_id, kelas_id, tanggal, status",             "Rekap kehadiran siswa"],
    ["Nilai",          "siswa_id, mapel_id, jenis, nilai_angka",          "Nilai akademik siswa"],
    ["Keuangan",       "siswa_id, jenis, jumlah, status",                 "Transaksi pembayaran SPP"],
    ["Pengumuman",     "judul, isi, kategori, prioritas",                 "Pengumuman dari guru/sekolah"],
    ["SuratEdaran",    "nomor_surat, judul, isi, lampiran",               "Surat resmi dari sekolah"],
    ["PPDB",           "no_pendaftaran, status, asal_sekolah",            "Data pendaftaran siswa baru"],
], col_widths=[3.0, 6.0, 7.5])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 7. Alur Pengguna (User Flow)
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. Alur Pengguna (User Flow)", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "7.1 Alur Login", level=2)
for txt in [
    "1. User membuka halaman login Teman Sekolah.",
    "2. User memasukkan username dan password.",
    "3. Sistem memvalidasi kredensial dan menghasilkan JWT token.",
    "4. Token disimpan di localStorage dan user diarahkan ke dashboard sesuai rolnya.",
    "5. Jika kredensial salah, sistem menampilkan pesan error.",
]:
    add_paragraph(doc, txt, size=11, space_after=3)

add_heading(doc, "7.2 Alur PPDB Online", level=2)
for txt in [
    "1. Calon siswa/orang tua mengakses formulir PPDB tanpa harus login.",
    "2. Calon siswa mengisi data diri, nilai rapor, dan dokumen pendukung.",
    "3. Sistem menyimpan data dan menampilkan nomor pendaftaran unik.",
    "4. TU mereview pendaftaran dan memperbarui status (Diterima/Ditolak/Pending).",
    "5. Calon siswa dapat mengecek status menggunakan nomor pendaftaran.",
    "6. Jika diterima, TU membuat akun siswa dan orang tua di sistem.",
]:
    add_paragraph(doc, txt, size=11, space_after=3)

add_heading(doc, "7.3 Alur Input Nilai Siswa", level=2)
for txt in [
    "1. Guru login dan membuka menu 'Input Nilai'.",
    "2. Guru memilih kelas, mata pelajaran, dan jenis penilaian (Harian/UTS/UAS).",
    "3. Sistem menampilkan daftar siswa di kelas tersebut.",
    "4. Guru mengisi nilai masing-masing siswa.",
    "5. Sistem menyimpan nilai dan menghitung rata-rata otomatis.",
    "6. Siswa dan orang tua dapat melihat nilai tersebut di dashboard mereka.",
]:
    add_paragraph(doc, txt, size=11, space_after=3)

add_heading(doc, "7.4 Alur Pembayaran SPP", level=2)
for txt in [
    "1. TU membuat tagihan SPP bulanan untuk seluruh siswa aktif.",
    "2. Orang tua menerima notifikasi tagihan di dashboard.",
    "3. Orang tua melakukan pembayaran via QRIS atau Virtual Account.",
    "4. TU mengkonfirmasi pembayaran dan mengubah status menjadi 'Lunas'.",
    "5. Rekap pembayaran tersedia di laporan keuangan Super Admin.",
]:
    add_paragraph(doc, txt, size=11, space_after=3)

# ─────────────────────────────────────────────────────────────────────────────
# 8. Desain Antarmuka (UI/UX)
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8. Desain Antarmuka (UI/UX)", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "8.1 Prinsip Desain", level=2)
add_bullet(doc, "Material Design: Menggunakan Material-UI (MUI) v5 sebagai sistem desain utama.")
add_bullet(doc, "Responsif: Layout menyesuaikan layar desktop, tablet, dan mobile (320px-2560px).")
add_bullet(doc, "Konsistensi: Komponen, warna, tipografi, dan spacing yang seragam di seluruh halaman.")
add_bullet(doc, "Aksesibilitas: Kontras warna memenuhi standar WCAG 2.1 Level AA.")
add_bullet(doc, "Efisiensi: Informasi kritis dapat ditemukan dalam maksimal 3 klik dari halaman mana pun.")

add_heading(doc, "8.2 Komponen Utama", level=2)
add_table(doc, ["Komponen", "Deskripsi"], [
    ["MainLayout",        "Layout utama dengan sidebar navigasi, header, dan area konten yang responsif"],
    ["ResponsiveTable",   "Tabel data yang dapat disortir, difilter, dan diekspor ke format Excel/PDF"],
    ["ResponsiveForm",    "Formulir input yang divalidasi secara real-time dengan pesan error informatif"],
    ["ResponsiveDialog",  "Modal dialog untuk konfirmasi aksi dan input data tambahan"],
    ["Dashboard Card",    "Kartu ringkasan statistik dengan warna, ikon, dan tren yang informatif"],
    ["Sidebar Navigation","Menu navigasi yang collapse di layar kecil dengan badge notifikasi"],
], col_widths=[4.5, 12.0])
doc.add_paragraph()

add_heading(doc, "8.3 Palet Warna", level=2)
add_table(doc, ["Warna", "Kode Hex", "Penggunaan"], [
    ["Primary Blue",   "#1565C0", "Warna utama brand, tombol CTA, header, ikon aktif"],
    ["Light Blue",     "#E3F2FD", "Background alternatif, highlight row tabel"],
    ["Success Green",  "#2E7D32", "Status sukses, konfirmasi sudah bayar, absensi Hadir"],
    ["Warning Orange", "#E65100", "Status perlu perhatian, tunggakan SPP"],
    ["Error Red",      "#C62828", "Validasi error, aksi berbahaya, absensi Alfa"],
    ["Neutral Gray",   "#616161", "Teks sekunder, placeholder, border, ikon nonaktif"],
], col_widths=[3.5, 3.0, 10.0])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 9. Keamanan
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "9. Keamanan", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "9.1 Mekanisme Autentikasi", level=2)
add_bullet(doc, "JWT (JSON Web Token): Token dengan algoritma HS256, expire 7 hari, disimpan di localStorage.")
add_bullet(doc, "Password Hashing: Semua password di-hash menggunakan Bcrypt dengan salt rounds minimum 10.")
add_bullet(doc, "Protected Routes: Setiap endpoint API diproteksi dengan middleware verifikasi JWT.")
add_bullet(doc, "Role-Based Access Control (RBAC): Akses dikontrol secara ketat berdasarkan peran pengguna.")

add_heading(doc, "9.2 Proteksi Data", level=2)
add_bullet(doc, "Input Validation: Validasi semua input di frontend dan backend untuk mencegah SQL Injection dan XSS.")
add_bullet(doc, "CORS Policy: Daftar domain yang diizinkan dikonfigurasi secara eksplisit di backend.")
add_bullet(doc, "Environment Variables: Semua konfigurasi sensitif (JWT secret, DB URL) disimpan di file .env.")
add_bullet(doc, "HTTPS: Wajib menggunakan SSL/TLS di lingkungan produksi.")
add_bullet(doc, "Audit Log: Setiap operasi kritis (hapus data, konfirmasi bayar) dicatat untuk audit trail.")

# ─────────────────────────────────────────────────────────────────────────────
# 10. Roadmap Pengembangan
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "10. Roadmap Pengembangan", level=1, color=(21, 101, 192))
add_section_line(doc)

add_table(doc, ["Fase", "Periode", "Fitur Utama", "Status"], [
    ["Fase 1 - MVP",       "Q1 2026", "Login/Auth, Dashboard semua role, Manajemen Siswa & Guru, Jadwal Pelajaran",     "Selesai (Dev)"],
    ["Fase 2 - Core",      "Q2 2026", "Absensi Digital, Input Nilai, E-Rapor, Keuangan SPP, PPDB Online",               "In Progress"],
    ["Fase 3 - Enhanced",  "Q3 2026", "Notifikasi Push, Pembayaran QRIS/VA, Laporan PDF, Surat Digital",                "Direncanakan"],
    ["Fase 4 - Scale",     "Q4 2026", "Multi-sekolah, Analytics Canggih, Mobile App Progressive Web App (PWA)",         "Direncanakan"],
    ["Fase 5 - Advanced",  "Q1 2027", "AI Rekomendasi Belajar, Integrasi Dapodik, SSO, Marketplace Konten Belajar",     "Visi Jangka Panjang"],
], col_widths=[3.0, 2.0, 9.0, 2.5])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 11. Metrik Keberhasilan
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "11. Metrik Keberhasilan (Success Metrics)", level=1, color=(21, 101, 192))
add_section_line(doc)

add_table(doc, ["Metrik", "Target", "Cara Ukur"], [
    ["Adoption Rate",           ">= 80% pengguna aktif dalam 3 bulan pertama",             "Analytics dashboard - login harian"],
    ["Task Completion Rate",    ">= 90% pengguna berhasil menyelesaikan tugas utama",       "User testing & log sistem"],
    ["Uptime",                  ">= 99.5% uptime bulanan",                                 "Monitoring server (Uptime Robot)"],
    ["Response Time API",       "<= 500ms untuk 95% request",                              "APM tools (New Relic / Datadog)"],
    ["Error Rate",              "<= 0.5% dari total request",                              "Logging sistem (Sentry)"],
    ["SPP Collection Rate",     "Peningkatan 20% tagihan terbayar bulan pertama",          "Laporan keuangan modul"],
    ["Teacher Efficiency",      "Pengurangan 50% waktu input absensi & nilai",             "Survey pengguna pasca-onboarding"],
    ["Parent Engagement",       ">= 70% orang tua aktif >= 2x/minggu",                    "Analytics dashboard aktifitas"],
    ["NPS Score",               ">= 50 (Net Promoter Score)",                              "Survey kepuasan pengguna berkala"],
    ["Digitalisasi Sekolah",    "100% sekolah partner gunakan sistem digital",             "Data laporan Super Admin"],
], col_widths=[4.5, 5.5, 6.5])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 12. Risiko & Mitigasi
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "12. Risiko & Mitigasi", level=1, color=(21, 101, 192))
add_section_line(doc)

add_table(doc, ["Risiko", "Dampak", "Probabilitas", "Strategi Mitigasi"], [
    ["Resistensi pengguna terhadap sistem baru", "Tinggi",        "Sedang", "Pelatihan onboarding, user guide, dan dukungan helpdesk aktif"],
    ["Kebocoran data siswa/keuangan",            "Sangat Tinggi", "Rendah", "Enkripsi data, audit keamanan berkala, backup harian otomatis"],
    ["Ketidakstabilan server saat ujian nasional","Tinggi",        "Sedang", "Auto-scaling, CDN, load balancing, maintenance terjadwal"],
    ["Perubahan regulasi pendidikan nasional",    "Sedang",        "Sedang", "Arsitektur modular untuk penyesuaian cepat tanpa gangguan besar"],
    ["Ketergantungan koneksi internet",           "Sedang",        "Tinggi", "Mode offline untuk absensi, sinkronisasi otomatis saat online"],
    ["Migrasi dari SQLite ke production DB",      "Sedang",        "Rendah", "ORM Sequelize mendukung multi-database, migrasi terjadwal"],
], col_widths=[4.5, 2.0, 2.0, 8.0])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 13. Batasan & Asumsi
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "13. Batasan & Asumsi", level=1, color=(21, 101, 192))
add_section_line(doc)

add_heading(doc, "13.1 Batasan Sistem", level=2)
add_bullet(doc, "Versi MVP menggunakan SQLite; migrasi ke MySQL/PostgreSQL wajib untuk produksi skala besar.")
add_bullet(doc, "Fitur pembayaran QRIS/VA memerlukan integrasi dengan payment gateway pihak ketiga (Midtrans/Xendit).")
add_bullet(doc, "Mobile App native (Android/iOS) tidak termasuk dalam ruang lingkup versi ini.")
add_bullet(doc, "Integrasi dengan sistem eksternal (Dapodik, BOS) direncanakan untuk fase berikutnya.")

add_heading(doc, "13.2 Asumsi", level=2)
add_bullet(doc, "Pengguna memiliki akses koneksi internet yang memadai.")
add_bullet(doc, "Setiap sekolah memiliki minimal satu administrator yang terlatih mengoperasikan sistem.")
add_bullet(doc, "Data siswa dan guru yang diinput diasumsikan valid dan telah diverifikasi pihak sekolah.")
add_bullet(doc, "Sistem dijalankan di browser modern dengan dukungan JavaScript ES6+.")

# ─────────────────────────────────────────────────────────────────────────────
# 14. Glosarium
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "14. Glosarium", level=1, color=(21, 101, 192))
add_section_line(doc)

add_table(doc, ["Istilah", "Definisi"], [
    ["SPA (Single Page Application)", "Aplikasi web yang memuat satu halaman HTML dan memperbarui konten secara dinamis"],
    ["RESTful API",                   "Antarmuka pemrograman yang menggunakan metode HTTP (GET, POST, PUT, DELETE)"],
    ["JWT Token",                     "Token terenkripsi yang digunakan untuk autentikasi stateless"],
    ["ORM",                           "Object Relational Mapping - teknik pemetaan objek program ke tabel database"],
    ["SQLite",                        "Database file-based ringan yang tidak memerlukan server terpisah"],
    ["CRUD",                          "Create, Read, Update, Delete - operasi dasar manajemen data"],
    ["RBAC",                          "Role-Based Access Control - kontrol akses berdasarkan peran pengguna"],
    ["E-Rapor",                       "Rapor elektronik digital yang dapat diakses secara online"],
    ["QRIS",                          "Quick Response Code Indonesian Standard - standar QR code pembayaran nasional"],
    ["Virtual Account (VA)",          "Nomor rekening virtual sementara untuk identifikasi pembayaran otomatis"],
    ["CDN",                           "Content Delivery Network - jaringan server untuk distribusi konten lebih cepat"],
    ["NPS",                           "Net Promoter Score - metrik kepuasan pelanggan (skala -100 sampai +100)"],
    ["PWA",                           "Progressive Web App - web app yang dapat diinstall seperti aplikasi native"],
], col_widths=[5.0, 11.5])
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# Riwayat Perubahan & Sign-off
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

rev_title = doc.add_paragraph()
rev_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = rev_title.add_run("Riwayat Perubahan Dokumen")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(21, 101, 192)

add_table(doc, ["Versi", "Tanggal", "Perubahan", "Penulis"], [
    ["1.0", "25 Maret 2026", "Pembuatan dokumen PRD awal berdasarkan hasil analisis project", "Tim Teman Sekolah"],
], col_widths=[1.5, 3.5, 9.0, 2.5])

doc.add_paragraph()
sign_off = doc.add_paragraph()
sign_off.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sign_off.add_run(
    "Dokumen ini adalah milik Tim Teman Sekolah dan bersifat rahasia.\n"
    "Teman Sekolah (c) 2026 - Sistem Manajemen Sekolah Terpadu")
run.font.size = Pt(10)
run.italic = True
run.font.color.rgb = RGBColor(120, 120, 120)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
output_path = r"d:\Project\Teman_Sekolah\PRD_Teman_Sekolah.docx"
doc.save(output_path)
print(f"PRD berhasil dibuat: {output_path}")
